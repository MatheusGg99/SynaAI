# tools.py
# System tools for file operations, web search, URL reading, image analysis, and system commands.

import os
import json
import re
import base64
import io
import subprocess
from ddgs import DDGS
from datetime import datetime
from PIL import Image
from core.constants import SYSTEM_APPS, SYSTEM_PATHS
from core.config import (
    MODEL, SYSTEM_PROMPT_FILE, MEMORY_FILE, PERSONALITY_FILE,
    TOOL_MODEL, PROJECT_ROOT
)
from core.utils import log_error, load_file, truncate_text
from core.llm import ask_syna, conversation_history
from core.clients import groq_client

# ------------------------------------------------------------------
# System prompt loader (cached)
# ------------------------------------------------------------------
_SYNA_SYSTEM_PROMPT = None

def _get_syna_system_prompt():
    """Loads Syna's system prompt once and caches it."""
    global _SYNA_SYSTEM_PROMPT
    if _SYNA_SYSTEM_PROMPT is None:
        from core.config import SYSTEM_PROMPT_FILE
        from core.utils import load_file
        _SYNA_SYSTEM_PROMPT = load_file(SYSTEM_PROMPT_FILE)
    return _SYNA_SYSTEM_PROMPT


# ------------------------------------------------------------------
# File reading
# ------------------------------------------------------------------
def read_file_content(path):
    """
    Reads content from text files, docx, pdf, or xlsx.
    Returns the content as a string, or an error message.
    """
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in ['.txt', '.md', '.py', '.gd', '.json', '.cfg', '.tscn', '.tres', '.import']:
            return load_file(path)
        elif ext == '.docx':
            import docx
            doc = docx.Document(path)
            full_text = []
            for p in doc.paragraphs:
                if p.text.strip():
                    full_text.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            return '\n'.join(full_text)
        elif ext == '.pdf':
            import PyPDF2
            text = ""
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
            return text
        elif ext in ['.xlsx', '.xls']:
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            sheet = wb.active
            data = []
            for row in sheet.iter_rows(values_only=True):
                data.append('\t'.join([str(cell) if cell else '' for cell in row]))
            return '\n'.join(data)
        else:
            return None
    except ImportError as e:
        lib = str(e).split("'")[1] if "'" in str(e) else "unknown"
        return f"[Error: library '{lib}' not installed. Run: pip install {lib}]"
    except Exception as e:
        return f"[Error reading file: {e}]"


# ------------------------------------------------------------------
# Web search and URL extraction
# ------------------------------------------------------------------
def get_search_urls(query, num_results=3):
    """Searches DuckDuckGo and returns a list of URLs."""
    try:
        results = []
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=num_results):
                results.append(r['href'])
        return results
    except ImportError:
        log_error("ddgs not installed.")
        print("⚠️ ddgs not installed. Web search will not work.")
        return []
    except Exception as e:
        log_error(f"Search error: {e}")
        print(f"⚠️ ddgs search error: {e}")
        return []


def extract_text_from_url(url, max_chars=5000):
    """
    Downloads a webpage and extracts clean text from paragraphs.
    Returns the text, truncated to max_chars if necessary.
    """
    try:
        import requests
        from bs4 import BeautifulSoup

        headers = {'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36'}
        response = requests.get(url, headers=headers, timeout=5)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, 'html.parser')
        # Remove scripts, styles, and common non-content sections
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()

        text = soup.get_text(separator='\n', strip=True)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        clean_text = '\n'.join(lines)

        if len(clean_text) > max_chars:
            clean_text = clean_text[:max_chars] + "... [truncated]"
        return clean_text
    except Exception as e:
        log_error(f"Error extracting text from {url}: {e}")
        return None


def summarize_web_content(content, user_query):
    """
    Summarizes the provided web content using Syna's personality.
    Returns a string with the summary.
    """
    if not content:
        return None
    try:
        system_prompt = _get_syna_system_prompt()
        # Add extra instructions to enforce personality and avoid generic phrases
        system_prompt += """

IMPORTANTE: Você é a Syna, parceira de Matheus. Responda com sua personalidade (tom sarcástico, usa apelidos como "Demandor", evite frases como "Olá! Estou aqui para ajudar..."). Seja direta, resuma o conteúdo em até 3 parágrafos, focando na pergunta do usuário. Não introduza com "Olá" ou "Estou aqui". Apenas vá direto ao ponto, como se estivesse conversando."""

        response = groq_client.chat.completions.create(
            model=TOOL_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Pergunta: {user_query}\n\nTexto da página:\n{content}"}
            ],
            max_tokens=500,
            temperature=0.7,
        )
        return response.choices[0].message.content
    except Exception as e:
        log_error(f"Error summarizing web content: {e}")
        return content[:1000] + "..."  # fallback


# ------------------------------------------------------------------
# System command detection and execution
# ------------------------------------------------------------------
def is_likely_system_command(user_input):
    """Quick pre-filter to avoid unnecessary API calls."""
    trigger_words = [
        "abrir", "abra", "abre", "open", "iniciar", "executar", "rodar",
        "criar arquivo", "cria arquivo", "novo arquivo", "escrever arquivo",
        "editar arquivo", "edita arquivo", "modificar arquivo",
        "mover arquivo", "move arquivo", "renomear",
        "deletar", "apagar", "excluir", "remover",
        "procurar arquivo", "buscar arquivo", "encontrar arquivo", "find",
        "pasta", "diretório", "folder",
        "terminal", "godot", "firefox", "spotify", "tlauncher",
        "saúde", "status", "sistema", "memória", "disco", "processos", "cpu", "uptime", "espaço", "ram"
    ]
    msg_lower = user_input.lower()
    return any(word in msg_lower for word in trigger_words)


def detect_command_local(user_input):
    """
    Detects system commands using local rules (no API calls).
    Returns a dict in the same format as detect_command, or None.
    """
    lower = user_input.lower()

    # System health
    health_triggers = ["status do sistema", "saúde do sistema", "diagnóstico do sistema", "status", "saúde"]
    if any(trigger in lower for trigger in health_triggers):
        return {"action": "system_health"}

    # App/folder keywords
    app_keywords = {
        "terminal": "terminal",
        "firefox": "firefox",
        "godot": "godot",
        "spotify": "spotify",
        "tlauncher": "tlauncher",
        "documentos": "documentos",
        "downloads": "downloads",
        "desktop": "desktop",
    }

    for keyword, target in app_keywords.items():
        if keyword in lower:
            if any(v in lower for v in ["abrir", "abre", "abra", "iniciar", "executar", "rodar"]):
                if target in ["documentos", "downloads", "desktop"]:
                    return {"action": "open_folder", "target": target}
                else:
                    return {"action": "open_app", "target": target}

    return None


def detect_command(user_input):
    """Uses the LLM to detect system commands from the user input."""
    lower = user_input.lower()
    health_triggers = ["status do sistema", "saúde do sistema", "diagnóstico do sistema", "status", "saúde"]
    if any(trigger in lower for trigger in health_triggers):
        return {"action": "system_health"}

    prompt = f"""Analyze the user's message and determine if it is a system command.
Message: "{user_input}"

Available apps: {list(SYSTEM_APPS.keys())}

Known folders — ALWAYS use these exact paths:
{json.dumps(SYSTEM_PATHS, ensure_ascii=False)}

If it is a system command, respond ONLY with JSON in this exact format:
{{"action": "open_app", "target": "app_name"}}
or
{{"action": "open_folder", "target": "folder_name"}}
or
{{"action": "find_file", "target": "file_name"}}
or
{{"action": "create_file", "target": "full/path/file.ext", "content": "actual file content"}}
or
{{"action": "edit_file", "target": "full/path/file.ext", "content": "new full content"}}
or
{{"action": "move_file", "target": "source_path", "destination": "dest_path"}}
or
{{"action": "delete_file", "target": "full/path/file.ext"}}
or
{{"action": "system_health"}}

Important rules:
- For folder paths, ALWAYS use the paths from the dictionary above. "downloads" = "~/Downloads", "documentos" = "~/Documents", "desktop" = "~/Desktop".
- If the user asks you to GENERATE or CREATE content (e.g., "write a sentence", "create a text", "put something creative"), you yourself must create that content and place it in the "content" field. Do not put literal instructions.
- Paths always use ~ for the user's home.
- If the user asks to GENERATE creative content, put "[GERAR] " followed by the instruction in the "content" field. Example: "[GERAR] write a motivational phrase with Syna's personality".
- If the user asks "status do sistema", "saúde do sistema" or similar, the action MUST be "system_health".

If it is NOT a system command, respond ONLY with:
{{"action": "none"}}

Respond ONLY with the JSON, no extra text."""

    try:
        response = groq_client.chat.completions.create(
            model=TOOL_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
            temperature=0.1,
            timeout=5,
        )
        raw = response.choices[0].message.content.strip()
        return json.loads(raw)
    except Exception:
        return {"action": "none"}


def get_system_health():
    """Collects basic system health information and returns a formatted string."""
    report = []
    report.append("🖥️ SISTEMA HEALTH REPORT\n")

    uptime = subprocess.getoutput("uptime -p").replace("up ", "")
    report.append(f"⏱️ Uptime: {uptime}")

    mem_used = subprocess.getoutput("free | grep '^Mem:' | awk '{print $3}'")
    mem_total = subprocess.getoutput("free | grep '^Mem:' | awk '{print $2}'")
    if mem_total and mem_used:
        percent = int(int(mem_used) / int(mem_total) * 100)
    else:
        percent = "?"
    mem_h = subprocess.getoutput("free -h | grep '^Mem:' | awk '{print $3 \"/\" $2}'")
    report.append(f"🧠 Memory: {mem_h} ({percent}%)")

    disk = subprocess.getoutput("df -h / | awk 'NR==2 {print $3 \"/\" $2 \" (\" $5 \")\"}'")
    report.append(f"💾 Disk (root): {disk}")

    load = subprocess.getoutput("cat /proc/loadavg | awk '{print $1 \", \" $2 \", \" $3}'")
    report.append(f"📊 Load average (1,5,15min): {load}")

    report.append("\n🔥 Top 5 processes by CPU:")
    top_cpu = subprocess.getoutput("ps aux --sort=-%cpu | head -6 | tail -5 | awk '{print $11 \" (\" $3 \"% CPU, \" $4 \"% MEM)\"}'")
    report.append(top_cpu if top_cpu else "No significant processes.")

    report.append("\n💧 Top 5 processes by memory:")
    top_mem = subprocess.getoutput("ps aux --sort=-%mem | head -6 | tail -5 | awk '{print $11 \" (\" $4 \"% MEM, \" $3 \"% CPU)\"}'")
    report.append(top_mem if top_mem else "No significant processes.")

    return "\n".join(report)


def execute_command(command):
    """Executes a system command based on the provided dictionary."""
    action = command.get("action")
    target = command.get("target", "").lower()

    # Security: validate that the app/folder exists in our lists
    if action == "open_app":
        app_found = False
        for app_name in SYSTEM_APPS:
            if app_name in target or target in app_name:
                app_found = True
                break
        if not app_found:
            log_error(f"Attempted to open unknown app: '{target}'")
            return None

    if action == "open_folder":
        folder_found = False
        for folder_name in SYSTEM_PATHS:
            if folder_name in target or target in folder_name:
                folder_found = True
                break
        if not folder_found:
            log_error(f"Attempted to open unknown folder: '{target}'")
            return None

    if action == "open_app":
        for app_name, app_path in SYSTEM_APPS.items():
            if app_name in target or target in app_name:
                expanded = os.path.expanduser(app_path)
                if os.path.exists(expanded):
                    if os.access(expanded, os.X_OK):
                        subprocess.Popen(f"'{expanded}'", shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    else:
                        subprocess.Popen(f"xdg-open '{expanded}'", shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                else:
                    subprocess.Popen(f"{app_path}", shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                return f"Abrindo {app_name}..."
        return f"App '{target}' não encontrado na lista."

    if action == "open_folder":
        for folder_name, folder_path in SYSTEM_PATHS.items():
            if folder_name in target or target in folder_name:
                expanded = os.path.expanduser(folder_path)
                if os.path.exists(expanded):
                    subprocess.Popen(f"xdg-open '{expanded}'", shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    return f"Abrindo pasta {folder_name}..."
                else:
                    return f"Pasta '{expanded}' não encontrada."
        return f"Pasta '{target}' não encontrada."

    if action == "find_file":
        result = os.popen(f"find ~ -name '*{target}*' 2>/dev/null | head -10").read()
        if result:
            return f"Encontrei:\n{result}"
        else:
            return f"Nenhum arquivo com '{target}' encontrado."

    if action == "create_file":
        target_path = command.get("target", "")
        content = command.get("content", "")
        expanded = os.path.expanduser(target_path)

        instrucoes = ["escreva", "crie", "gere", "coloque", "faça", "uma frase", "um texto", "algo"]
        precisa_gerar = any(palavra in content.lower() for palavra in instrucoes)

        if precisa_gerar:
            try:
                gen_response = groq_client.chat.completions.create(
                    model=TOOL_MODEL,
                    messages=[
                        {
                            "role": "system",
                            "content": "Você é Syna. Gere o conteúdo solicitado com sua personalidade. Responda APENAS com o conteúdo final, sem explicações, sem prefácios."
                        },
                        {
                            "role": "user",
                            "content": content
                        }
                    ],
                    max_tokens=1000,
                    temperature=0.9,
                )
                content = gen_response.choices[0].message.content.strip()
            except Exception as e:
                log_error(e)
                content = f"Erro ao gerar conteúdo: {e}"

        try:
            dir_path = os.path.dirname(expanded)
            if dir_path:
                os.makedirs(dir_path, exist_ok=True)
            with open(expanded, "w", encoding="utf-8") as f:
                f.write(content)
            return f"Arquivo criado: {expanded}\n\nConteúdo:\n{content}"
        except Exception as e:
            log_error(e)
            return f"Erro ao criar arquivo: {e}"

    if action == "edit_file":
        target_path = command.get("target", "")
        content = command.get("content", "")
        expanded = os.path.expanduser(target_path)

        instrucoes = ["escreva", "crie", "gere", "coloque", "faça", "uma frase", "um texto", "algo", "adicione"]
        precisa_gerar = any(palavra in content.lower() for palavra in instrucoes)

        if precisa_gerar:
            try:
                current_content = load_file(expanded)
                gen_response = groq_client.chat.completions.create(
                    model=TOOL_MODEL,
                    messages=[
                        {
                            "role": "system",
                            "content": "Você é Syna. Gere o conteúdo solicitado com sua personalidade. Responda APENAS com o conteúdo final, sem explicações, sem prefácios."
                        },
                        {
                            "role": "user",
                            "content": f"Conteúdo atual do arquivo:\n{current_content}\n\nInstrução: {content}"
                        }
                    ],
                    max_tokens=1000,
                    temperature=0.9,
                )
                content = gen_response.choices[0].message.content.strip()
            except Exception as e:
                log_error(e)
                content = f"Erro ao gerar conteúdo: {e}"

        if os.path.exists(expanded):
            try:
                with open(expanded, "w", encoding="utf-8") as f:
                    f.write(content)
                return f"Arquivo editado: {expanded}\n\nNovo conteúdo:\n{content}"
            except Exception as e:
                log_error(e)
                return f"Erro ao editar arquivo: {e}"
        else:
            return f"Arquivo não encontrado: {expanded}"

    if action == "move_file":
        target_path = command.get("target", "")
        destination = command.get("destination", "")
        expanded_src = os.path.expanduser(target_path)
        expanded_dst = os.path.expanduser(destination)
        if os.path.exists(expanded_src):
            try:
                import shutil
                shutil.move(expanded_src, expanded_dst)
                return f"Arquivo movido para: {expanded_dst}"
            except Exception as e:
                log_error(e)
                return f"Erro ao mover arquivo: {e}"
        else:
            return f"Arquivo não encontrado: {expanded_src}"

    if action == "delete_file":
        target_path = command.get("target", "")
        expanded = os.path.expanduser(target_path)
        if os.path.exists(expanded):
            try:
                os.remove(expanded)
                return f"Arquivo deletado: {expanded}"
            except Exception as e:
                log_error(e)
                return f"Erro ao deletar arquivo: {e}"
        else:
            return f"Arquivo não encontrado: {expanded}"

    if action == "system_health":
        report = get_system_health()
        return f"{report}\n\n💬 O que você gostaria de saber mais? Posso analisar algum ponto específico (memória, disco, processos) se quiser."

    return None


def handle_system_command(user_input):
    """Handles system commands if the input matches local triggers."""
    if not is_likely_system_command(user_input):
        return None

    command = detect_command_local(user_input)
    if command:
        return execute_command(command)

    command = detect_command(user_input)
    if command.get("action") == "none":
        return None

    return execute_command(command)


# ------------------------------------------------------------------
# Web search (public interface)
# ------------------------------------------------------------------
def handle_web_search(user_input):
    """
    Detects a web search query, fetches content from top results,
    and returns a synthesized summary with Syna's personality.
    """
    web_triggers = [
        "pesquise", "pesquisa", "procure por", "busque por", "procura na internet",
        "busca na web", "pesquisar", "buscar", "procure", "busque"
    ]
    lower = user_input.lower()

    for trigger in web_triggers:
        if trigger in lower:
            idx = lower.find(trigger)
            query = user_input[idx + len(trigger):].strip()
            query = re.sub(r'^[\s,:.]+', '', query)
            query = re.sub(r'^(sobre|por|acerca de|a respeito de)\s+', '', query, flags=re.IGNORECASE)

            if not query:
                return None

            urls = get_search_urls(query, num_results=2)
            if not urls:
                return f"Não consegui encontrar resultados relevantes para '{query}'."

            summaries = []
            for url in urls:
                content = extract_text_from_url(url, max_chars=4000)
                if content:
                    summary = summarize_web_content(content, query)
                    if summary:
                        summaries.append(f"**Fonte:** {url}\n{summary}")

            if not summaries:
                return f"Até encontrei algumas páginas sobre '{query}', mas não consegui extrair informações úteis delas. 🙃"

            combined = "\n\n---\n\n".join(summaries)
            return f"🔍 Pesquisei sobre '{query}' e analisei os conteúdos:\n\n{combined}\n\n💬 Com base nisso, o que você quer saber exatamente? Posso dar minha opinião se quiser."

    return None


# ------------------------------------------------------------------
# Local file and image reading (user-facing)
# ------------------------------------------------------------------
def try_read_file(user_input):
    """Attempts to read a local file mentioned in the user input."""
    patterns = [
        r'l[eê]ia?\s+(?:o\s+)?arquivo\s+(\S+)',
        r'abra?\s+(?:o\s+)?arquivo\s+(\S+)',
        r'analise?\s+(?:o\s+)?arquivo\s+(\S+)',
        r'veja?\s+(?:o\s+)?arquivo\s+(\S+)',
        r'leia\s+(\S+\.(?:gd|py|txt|md|json|cfg|tscn|tres|import))',
    ]
    for pattern in patterns:
        match = re.search(pattern, user_input, re.IGNORECASE)
        if match:
            path = match.group(1)
            expanded = os.path.expanduser(path)
            if os.path.exists(expanded):
                content = read_file_content(expanded)
                if content:
                    if len(content) > 8000:
                        content = content[:8000] + "\n\n[arquivo truncado]"
                    return f"[Arquivo carregado: {expanded}]\n\n{content}"
                else:
                    return f"[Arquivo vazio: {expanded}]"
            else:
                return f"[Arquivo não encontrado: {expanded}]"
    return None


def try_read_image(user_input):
    """Attempts to read an image file mentioned in the user input and returns a description."""
    patterns = [
        r'(?:analise|descreva|veja|olhe|leia|o\s+que\s+(?:tem|há))\s+(?:a\s+)?(?:imagem|foto|figura|print)\s+(?:do\s+)?(?:arquivo\s+)?[\'"]?([^\'"]+\.(?:png|jpg|jpeg|gif|bmp|webp))[\'"]?',
        r'(?:imagem|foto|figura)\s+[\'"]?([^\'"]+\.(?:png|jpg|jpeg|gif|bmp|webp))[\'"]?',
    ]
    for pattern in patterns:
        match = re.search(pattern, user_input, re.IGNORECASE)
        if match:
            path = match.group(1).strip()
            expanded = os.path.expanduser(path)
            if not os.path.exists(expanded):
                return f"Arquivo não encontrado: {expanded}"
            try:
                img = Image.open(expanded)
                img.thumbnail((1024, 1024))

                if img.mode in ('RGBA', 'LA', 'P'):
                    rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                    rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = rgb_img

                buffer = io.BytesIO()
                img.save(buffer, format="JPEG", quality=85)
                img_base64 = base64.b64encode(buffer.getvalue()).decode()

                vision_response = groq_client.chat.completions.create(
                    model="meta-llama/llama-4-scout-17b-16e-instruct",
                    messages=[
                        {"role": "system", "content": load_file(SYSTEM_PROMPT_FILE)},
                        {
                            "role": "user",
                            "content": [
                                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_base64}"}},
                                {"type": "text", "text": f"O usuário pediu: '{user_input}'. Analise a imagem do arquivo '{path}' e responda com sua personalidade."}
                            ]
                        }
                    ],
                    max_tokens=1024,
                    temperature=0.7,
                )
                reply = vision_response.choices[0].message.content
                conversation_history.append({
                    "role": "user",
                    "content": f"[Imagem analisada: {path}] {user_input}"
                })
                conversation_history.append({
                    "role": "assistant",
                    "content": reply
                })
                return reply
            except Exception as e:
                log_error(f"Error processing image {expanded}: {e}")
                return f"Erro ao processar imagem: {e}"
    return None


# ------------------------------------------------------------------
# Screenshot capture
# ------------------------------------------------------------------
def take_screenshot(user_input):
    """Captures the screen if the user asks for it, and returns an analysis."""
    keywords = ["olha minha tela", "veja minha tela", "screenshot", "print da tela", "o que está na tela"]
    lower = user_input.lower()
    if not any(k in lower for k in keywords):
        return None

    try:
        import mss
        import mss.tools
        from PIL import Image
        import base64
        import io

        screenshot_path = os.path.join(PROJECT_ROOT, "assets", "screenshot.png")

        with mss.mss() as sct:
            monitor = sct.monitors[1]
            screenshot = sct.grab(monitor)
            mss.tools.to_png(screenshot.rgb, screenshot.size, output=screenshot_path)

        img = Image.open(screenshot_path)
        max_width = 1920
        max_height = 1080
        img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)

        buffer = io.BytesIO()
        img.save(buffer, format="JPEG", quality=90, optimize=True)
        img_base64 = base64.b64encode(buffer.getvalue()).decode()

        vision_response = groq_client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=[
                {
                    "role": "system",
                    "content": load_file(SYSTEM_PROMPT_FILE)
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{img_base64}"
                            }
                        },
                        {
                            "type": "text",
                            "text": f"""O usuário pediu: '{user_input}'. Você é a Syna, com acesso ao sistema e capacidade de análise técnica. 
Com base APENAS no que está visível na imagem, forneça uma resposta que atenda ao pedido do usuário.
Se ele pediu para opinar sobre código, analise o código visível.
Se pediu para identificar algo, identifique.
Se pediu uma ação (ex: "abre o terminal"), diga que vê o terminal e execute o comando internamente.
Não apenas descreva a tela — aja conforme o pedido, dentro das suas capacidades."""
                        }
                    ]
                }
            ],
            max_tokens=1024,
            temperature=0.1,
        )

        vision_reply = vision_response.choices[0].message.content

        conversation_history.append({
            "role": "user",
            "content": f"[Analisei minha tela] {user_input}"
        })
        conversation_history.append({
            "role": "assistant",
            "content": vision_reply
        })

        return vision_reply

    except Exception as e:
        log_error(e)
        return f"Erro ao capturar tela: {e}"


# ------------------------------------------------------------------
# Desktop notifications
# ------------------------------------------------------------------
def send_desktop_notification(title, message, duration_ms=None):
    """Sends a native desktop notification with Syna's icon."""
    try:
        import subprocess
        icon_path = os.path.join(PROJECT_ROOT, "assets", "icon.jpg")
        cmd = ['notify-send', title, message, '--icon', icon_path]
        if duration_ms:
            cmd += ['-t', str(duration_ms)]
        subprocess.run(cmd, check=True)
        return True
    except Exception as e:
        log_error(f"Error sending notification: {e}")
        return False


# ------------------------------------------------------------------
# URL reading (public interface)
# ------------------------------------------------------------------
def handle_url(user_input):
    """
    Detects a URL in the user input, extracts its content,
    injects it into the conversation history, and returns a summary
    with Syna's personality.
    """
    url_match = re.search(r'(https?://\S+)', user_input)
    if not url_match:
        return None

    url = url_match.group(1).strip()
    query = user_input.replace(url, '').strip()
    if not query:
        query = "Resuma o conteúdo desta página."

    content = extract_text_from_url(url, max_chars=5000)
    if not content:
        return f"❌ Não consegui acessar {url}. Verifique o link."

    # Inject into history for future reference
    conversation_history.append({
        "role": "user",
        "content": f"[Conteúdo da página {url} carregado para análise]"
    })
    conversation_history.append({
        "role": "assistant",
        "content": f"Texto da página (truncado):\n{content[:2000]}..."
    })

    summary = summarize_web_content(content, query)
    return f"🔗 {summary}"